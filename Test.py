from docx.enum.text import WD_COLOR_INDEX
from docx import Document

# doc = docx.Document()
# p = doc.add_paragraph("Hello")
# p.add_run("Write more things in this document")
# p.add_run("Highlight this in Red").font.highlight_color = WD_COLOR_INDEX.RED
# doc.save("Gibbrish.docx")
import docx
import nltk
from tkinter import *
from tkinter import messagebox

# reading the list of names of students present in the class
file1 = open("list.txt", "r")
text1 = file1.readlines()
str1 = ''.join(text1)
str = str1.split(' ')
print(str)  # str renders each name


def getme():
    if(flag == 1):
        messagebox.showinfo("Result", "THE DOCUMENT IS PLAGIARISED")
    else:
        messagebox.showinfo("Result", "THE DOCUMENT IS NOT PLAGIARISED")


# returns each word file in form of a list separated by full stops
def ReadingTextDocuments(testPython):
    doc = docx.Document(testPython)

    completedText = []
    for paragraph in doc.paragraphs:
        completedText.append(paragraph.text)

    s = '\n'.join(completedText)
    s1 = s.split('.')
    return s1


def highlightme(person):
    doc = docx.Document(person)
    p = doc.add_paragraph("This Document has been checked for Plagiarism")
    for line1 in doc.paragraphs:
        for line2 in allofit:
            if(line1 == line2):
                print(line1)

                tempporarystorage = line1

                print(tempporarystorage.text)
                p.add_run(
                    tempporarystorage.text).font.highlight_color = WD_COLOR_INDEX.RED
                doc.save(person)


name = input("Enter the name of the first document\t")
name1 = name+'.docx'
flag = 0
doc1 = ReadingTextDocuments(name1)
doc1.pop()  # Reads the user input file, splits it into a list with . as a delimiter
LOC = len(doc1)
allofit = []
for out in str:
    name2 = out+'.docx'

    doc2 = ReadingTextDocuments(name2)
    doc2.pop()
    cnt1 = 0
    for i in doc1:
        print(i)
        for j in doc2:
           # print(i,j)
            if(i == j):
                print(i)
                allofit.append(j)
                cnt1 += 1
    # print(cnt1,LOC)
    per = (cnt1/LOC)*100
    if(per > 55 and out != name):
        flag = 1
    if(out != name):
        print("The percentage of document that has been copied from ",
              out, "is: ", per, "% \n")
print("\n\n\n\n\n\n\n")
print("this is all of it \n.")
print(allofit)
highlightme(name1)
getme()


# doc.text = ''
# a = doc.add_paragraph("This is checked for plagiarism")
# a.add_run(i).font.highlight_color = WD_COLOR_INDEX.RED


root = Tk()
button = Button(root, command=getme)
# button.pack()
# root.geometry("400x400")
# root.mainloop()

# doc = docx.Document('Sanyam.docx')
# p = doc.add_paragraph("This Document has been checked for Plagiarism")
# for i in doc.paragraphs:
#     tempporarystorage = i
#     i = ""
#     print(tempporarystorage.text)
#     p.add_run(tempporarystorage.text).font.highlight_color = WD_COLOR_INDEX.RED
#     doc.save("Gibbrish.docx")
